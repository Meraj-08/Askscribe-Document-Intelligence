import os
import logging
from typing import List, Optional
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
import io

# Simple text splitter implementation
class SimpleTextSplitter:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = text[:end].rfind('.')
                last_newline = text[:end].rfind('\n')
                
                break_point = max(last_period, last_newline)
                if break_point > start + self.chunk_size // 2:
                    end = break_point + 1
            
            chunk = text[start:end].strip()
            if len(chunk) > 50:  # Only include meaningful chunks
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks

class DocumentProcessor:
    """Handles document text extraction and chunking"""
    
    def __init__(self):
        self.text_splitter = SimpleTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document based on file type"""
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            elif file_type == 'txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logging.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with OCR fallback for scanned documents"""
        try:
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # First try to extract text directly
                page_text = page.get_text()
                
                # If very little text found, use OCR
                if len(page_text.strip()) < 50:
                    logging.info(f"Page {page_num + 1} has minimal text, using OCR")
                    
                    # Get page as image
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Use OCR to extract text
                    ocr_text = pytesseract.image_to_string(image)
                    text += f"\n--- Page {page_num + 1} (OCR) ---\n{ocr_text}\n"
                else:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            doc.close()
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logging.error(f"TXT processing error: {e}")
            raise
    
    def create_chunks(self, text: str) -> List[str]:
        """Split text into chunks for embedding"""
        try:
            # Clean and preprocess text
            text = self._preprocess_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Filter out very short chunks
            chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
            
            logging.info(f"Created {len(chunks)} chunks from text")
            return chunks
            
        except Exception as e:
            logging.error(f"Text chunking error: {e}")
            raise
    
    def _preprocess_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove special characters that might interfere
        text = text.replace('\x00', '')  # Remove null characters
        
        return text
    
    def get_document_info(self, file_path: str, file_type: str) -> dict:
        """Get basic information about the document"""
        info = {
            'pages': 1,
            'word_count': 0,
            'char_count': 0
        }
        
        try:
            # Extract text to count words and characters
            text = self.extract_text(file_path, file_type)
            info['word_count'] = len(text.split())
            info['char_count'] = len(text)
            
        except Exception as e:
            logging.error(f"Document info extraction error: {e}")
        
        return info
